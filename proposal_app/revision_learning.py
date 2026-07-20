from __future__ import annotations

from collections import Counter
from datetime import datetime, timezone
from difflib import SequenceMatcher
from hashlib import sha256
from io import BytesIO
import re
from typing import Iterator

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph

from .models import RevisionAnalysis


PROPOSAL_NUMBER_RE = re.compile(r"P0\d{2}-[0-9A-Za-z]+", re.IGNORECASE)
TOKEN_RE = re.compile(r"[a-z0-9]+")
SAFE_KEY_RE = re.compile(r"[^a-z0-9]+")


def iter_blocks(document: DocumentType) -> Iterator[Paragraph | Table]:
    for child in document.element.body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, document)
        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


def extract_docx_lines(data: bytes) -> list[str]:
    document = Document(BytesIO(data))
    lines: list[str] = []
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = " ".join(block.text.split())
            if text:
                lines.append(text)
        else:
            for row in block.rows:
                cells = [" ".join(cell.text.split()) for cell in row.cells]
                text = " | ".join(cells).strip(" |")
                if text:
                    lines.append(text)
    return lines


def compare_docx(draft_bytes: bytes, final_bytes: bytes, max_blocks: int = 80) -> dict:
    draft_lines = extract_docx_lines(draft_bytes)
    final_lines = extract_docx_lines(final_bytes)
    matcher = SequenceMatcher(None, draft_lines, final_lines, autojunk=False)
    differences: list[dict] = []
    for operation, i1, i2, j1, j2 in matcher.get_opcodes():
        if operation == "equal":
            continue
        differences.append(
            {
                "operation": operation,
                "draft": "\n".join(draft_lines[i1:i2])[:4000],
                "final": "\n".join(final_lines[j1:j2])[:4000],
            }
        )
        if len(differences) >= max_blocks:
            break
    return {
        "similarity": round(matcher.ratio(), 4),
        "draft_line_count": len(draft_lines),
        "final_line_count": len(final_lines),
        "differences": differences,
        "draft_sha256": sha256(draft_bytes).hexdigest(),
        "final_sha256": sha256(final_bytes).hexdigest(),
    }


def normalize_candidate_key(value: str) -> str:
    return SAFE_KEY_RE.sub("-", value.lower()).strip("-")[:80]


def revision_record(
    proposal_number: str,
    final_filename: str,
    repository_path: str,
    comparison: dict,
    analysis: RevisionAnalysis,
) -> dict:
    candidates = []
    for candidate in analysis.candidates:
        key = normalize_candidate_key(candidate.key)
        if not key:
            continue
        candidates.append({**candidate.model_dump(mode="json"), "key": key})
    return {
        "id": comparison["final_sha256"][:16],
        "created_at": datetime.now(timezone.utc).isoformat(),
        "proposal_number": proposal_number.upper(),
        "final_filename": final_filename,
        "repository_path": repository_path,
        "summary": analysis.summary,
        "similarity": comparison["similarity"],
        "draft_sha256": comparison["draft_sha256"],
        "final_sha256": comparison["final_sha256"],
        "differences": comparison["differences"],
        "candidates": candidates,
    }


def aggregate_rule_candidates(state: dict, minimum_occurrences: int = 3) -> list[dict]:
    approved_keys = {rule["key"] for rule in state.get("approved_rules", [])}
    occurrences: Counter[str] = Counter()
    candidate_details: dict[str, dict] = {}

    for record in state.get("records", []):
        seen_in_record: set[str] = set()
        for candidate in record.get("candidates", []):
            key = normalize_candidate_key(candidate.get("key", ""))
            if not key or not candidate.get("reusable", False) or key in approved_keys:
                continue
            candidate_details[key] = candidate
            if key not in seen_in_record:
                occurrences[key] += 1
                seen_in_record.add(key)

    ready = []
    for key, count in occurrences.items():
        if count >= minimum_occurrences:
            ready.append({**candidate_details[key], "key": key, "occurrences": count})
    return sorted(ready, key=lambda item: (-item["occurrences"], item["key"]))


def classify_project(filename: str, introduction: str) -> str:
    source = f"{filename} {introduction}".lower()
    rules = [
        ("interior slab / hand auger", ("hand auger", "slab movement", "slab investigation")),
        ("foundation restoration / bearing review", ("foundation restoration", "bearing review")),
        ("retaining wall", ("retaining wall",)),
        ("single-family dwelling", ("single-family", "single family", "dwelling")),
        ("townhouse / residential development", ("townhouse",)),
        ("building expansion", ("expansion", "maintenance garage", "clubhouse")),
        ("small building", ("pole barn",)),
        ("industrial / heavy-use site", ("industrial", "precision drilling", "ward tire")),
        ("roadway / pavement", ("roadway", "asphalt ramp", "112th street")),
        ("commercial building / site development", ("proposed building", "commercial building")),
    ]
    for label, needles in rules:
        if any(needle in source for needle in needles):
            return label
    return "general building / site development"


def classify_methods(text: str) -> list[str]:
    lowered = text.lower()
    tests = [
        ("drilling", ("borehole", "drilling rig", "solid stem auger")),
        ("test pits", ("test pit", "test-pit", "excavator")),
        ("hand auger", ("hand auger", "hand-auger")),
        ("groundwater monitoring", ("piezometer", "groundwater reading")),
        ("coring", ("core drill", "coring")),
        ("laboratory testing", ("laboratory testing", "moisture content", "atterberg")),
    ]
    return [label for label, needles in tests if any(needle in lowered for needle in needles)]


def proposal_index_record(data: bytes, filename: str) -> dict:
    document = Document(BytesIO(data))
    sections: dict[str, list[str]] = {"Front matter": []}
    current = "Front matter"
    headings: list[str] = []
    for block in iter_blocks(document):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if not text:
                continue
            if block.style and block.style.name.startswith("Heading"):
                current = text
                headings.append(text)
                sections.setdefault(current, [])
            else:
                sections.setdefault(current, []).append(text)
        else:
            rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in block.rows]
            sections.setdefault(current, []).append("\n".join(rows))

    flattened = {heading: "\n".join(lines).strip() for heading, lines in sections.items()}
    full_text = "\n".join(flattened.values())
    number_match = PROPOSAL_NUMBER_RE.search(filename) or PROPOSAL_NUMBER_RE.search(full_text)
    proposal_number = number_match.group(0).upper() if number_match else ""
    tokens = sorted(
        {token for token in TOKEN_RE.findall(f"{filename} {full_text}".lower()) if len(token) > 2}
    )
    return {
        "filename": filename,
        "proposal_number": proposal_number,
        "project_type": classify_project(filename, flattened.get("Introduction", "") or full_text),
        "methods": classify_methods(flattened.get("Geotechnical Field Program", full_text)),
        "has_cost_table": any(len(table.columns) == 6 for table in document.tables),
        "headings": headings,
        "sections": flattened,
        "tokens": tokens,
    }
