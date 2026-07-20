from __future__ import annotations

import base64
import csv
import io
import mimetypes
from dataclasses import dataclass
from email import policy
from email.parser import BytesParser
from pathlib import Path

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


MAX_FILE_BYTES = 20 * 1024 * 1024


@dataclass(frozen=True)
class UploadedDocument:
    name: str
    data: bytes
    mime_type: str = ""

    @property
    def suffix(self) -> str:
        return Path(self.name).suffix.lower()

    @property
    def mime(self) -> str:
        return self.mime_type or mimetypes.guess_type(self.name)[0] or "application/octet-stream"


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8-sig", "utf-8", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_eml(document: UploadedDocument) -> tuple[str, list[UploadedDocument]]:
    message = BytesParser(policy=policy.default).parsebytes(document.data)
    lines = [
        f"From: {message.get('From', '')}",
        f"To: {message.get('To', '')}",
        f"Cc: {message.get('Cc', '')}",
        f"Subject: {message.get('Subject', '')}",
        f"Date: {message.get('Date', '')}",
        "",
    ]
    attachments: list[UploadedDocument] = []
    body_found = False
    for part in message.walk():
        disposition = part.get_content_disposition()
        filename = part.get_filename()
        content_type = part.get_content_type()
        if disposition == "attachment" or filename:
            payload = part.get_payload(decode=True) or b""
            if payload and len(payload) <= MAX_FILE_BYTES:
                attachments.append(
                    UploadedDocument(filename or "attachment", payload, content_type)
                )
            continue
        if not body_found and content_type == "text/plain":
            try:
                lines.append(part.get_content())
            except Exception:
                lines.append(decode_text(part.get_payload(decode=True) or b""))
            body_found = True
    if not body_found:
        try:
            lines.append(message.get_body(preferencelist=("plain", "html")).get_content())
        except Exception:
            pass
    return "\n".join(lines).strip(), attachments


def extract_docx(data: bytes) -> str:
    document = Document(io.BytesIO(data))
    parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text.replace("\n", " | ") for cell in row.cells))
    return "\n".join(parts)


def extract_xlsx(data: bytes) -> str:
    workbook = load_workbook(io.BytesIO(data), data_only=True, read_only=True)
    output: list[str] = []
    for sheet in workbook.worksheets:
        output.append(f"Worksheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(value.strip() for value in values):
                output.append("\t".join(values))
    return "\n".join(output)


def extract_pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages).strip()


def extract_text(document: UploadedDocument) -> tuple[str, list[UploadedDocument]]:
    suffix = document.suffix
    if suffix == ".eml":
        return extract_eml(document)
    if suffix == ".docx":
        return extract_docx(document.data), []
    if suffix in {".xlsx", ".xlsm"}:
        return extract_xlsx(document.data), []
    if suffix == ".pdf":
        return extract_pdf_text(document.data), []
    if suffix in {".csv", ".tsv"}:
        text = decode_text(document.data)
        delimiter = "\t" if suffix == ".tsv" else ","
        rows = csv.reader(io.StringIO(text), delimiter=delimiter)
        return "\n".join("\t".join(row) for row in rows), []
    if suffix in {".txt", ".md", ".html", ".htm"}:
        return decode_text(document.data), []
    return "", []


def expand_documents(documents: list[UploadedDocument]) -> tuple[list[str], list[UploadedDocument]]:
    text_parts: list[str] = []
    binary_documents: list[UploadedDocument] = []
    queue = list(documents)
    while queue:
        document = queue.pop(0)
        if len(document.data) > MAX_FILE_BYTES:
            text_parts.append(f"[{document.name} was skipped because it exceeds 20 MB]")
            continue
        extracted, attachments = extract_text(document)
        if extracted:
            text_parts.append(f"--- {document.name} ---\n{extracted}")
        queue.extend(attachments)
        if document.suffix == ".pdf" or document.mime.startswith("image/"):
            binary_documents.append(document)
    return text_parts, binary_documents


def build_openai_content(
    email_text: str,
    notes: str,
    documents: list[UploadedDocument],
) -> list[dict]:
    extracted_text, binary_documents = expand_documents(documents)
    text = "\n\n".join(
        part
        for part in [
            "USER-PASTED REQUEST EMAIL:\n" + email_text.strip() if email_text.strip() else "",
            "USER NOTES:\n" + notes.strip() if notes.strip() else "",
            *extracted_text,
        ]
        if part
    )
    content: list[dict] = [{"type": "input_text", "text": text or "No request text supplied."}]
    for document in binary_documents:
        encoded = base64.b64encode(document.data).decode("ascii")
        if document.suffix == ".pdf":
            content.append(
                {
                    "type": "input_file",
                    "filename": document.name,
                    "file_data": f"data:application/pdf;base64,{encoded}",
                }
            )
        elif document.mime.startswith("image/"):
            content.append(
                {
                    "type": "input_image",
                    "image_url": f"data:{document.mime};base64,{encoded}",
                }
            )
    return content

