from __future__ import annotations

import io
import re
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

from docx import Document

from .costs import as_number
from .models import ProposalFacts


PROPOSAL_NUMBER_PATTERN = re.compile(r"P0\d{2}-[0-9A-Za-z]+", re.IGNORECASE)
MONEY_PATTERN = re.compile(r"\$\s*([\d,]+(?:\.\d{2})?)")


@dataclass(frozen=True)
class UploadedProposalDetails:
    facts: ProposalFacts
    work_authorization_scope: str
    budget: float


def client_email_from_proposal(facts: ProposalFacts) -> tuple[str, str]:
    subject_parts = [facts.proposal_number, facts.project_name, "Geotechnical Proposal"]
    subject = " - ".join(part.strip() for part in subject_parts if part.strip())
    first_name = facts.contact_name.strip().split()[0] if facts.contact_name.strip() else "there"
    lines = []
    if facts.contact_email.strip():
        lines.extend([f"To: {facts.contact_email.strip()}", ""])
    lines.extend(
        [
            f"Hi {first_name},",
            "",
            (
                f"Please find attached our geotechnical proposal for {facts.project_name}."
                if facts.project_name
                else "Please find attached our geotechnical proposal."
            ),
            "",
            "Please sign and return the Work Authorization form to proceed with the work.",
            "",
            "Regards,",
            "Steven",
        ]
    )
    return subject, "\n".join(lines)


def _nonempty_lines(value: str) -> list[str]:
    return [line.strip() for line in value.splitlines() if line.strip()]


def _money(value: str) -> float | None:
    matches = MONEY_PATTERN.findall(value)
    return as_number(matches[-1]) if matches else None


def _iso_date(value: str) -> str:
    for date_format in ("%B %d, %Y", "%b %d, %Y", "%Y-%m-%d"):
        try:
            return datetime.strptime(value.strip(), date_format).date().isoformat()
        except ValueError:
            continue
    return ""


def _proposal_budget(document) -> float | None:
    for table in document.tables:
        for row in reversed(table.rows):
            values = [cell.text.strip() for cell in row.cells]
            if any("estimated cost" in value.casefold() for value in values):
                amount = _money(" ".join(values))
                if amount is not None:
                    return amount
    for paragraph in document.paragraphs:
        if "total estimated" in paragraph.text.casefold():
            amount = _money(paragraph.text)
            if amount is not None:
                return amount
    return None


def _project_details(front_lines: list[str]) -> tuple[str, str]:
    re_index = next(
        (
            index
            for index, line in enumerate(front_lines)
            if re.match(r"^re\s*:", line, flags=re.IGNORECASE)
        ),
        None,
    )
    if re_index is None:
        return "", ""
    candidate_lines: list[str] = []
    for index, line in enumerate(front_lines[re_index:], start=re_index):
        if index == re_index:
            line = re.sub(r"^re\s*:\s*", "", line, flags=re.IGNORECASE).strip()
        if line and "request for proposal" not in line.casefold():
            candidate_lines.append(line)
    if not candidate_lines:
        return "", ""
    return candidate_lines[0], "\n".join(candidate_lines[1:])


def _project_from_identifier(
    paragraphs,
    proposal_number: str,
    client_name: str,
) -> str:
    for paragraph in reversed(paragraphs):
        for line in reversed(_nonempty_lines(paragraph.text)):
            if not line.casefold().startswith(proposal_number.casefold()):
                continue
            candidate = line[len(proposal_number) :].strip(" -")
            client_suffix = f" - {client_name}" if client_name else ""
            if client_suffix and candidate.casefold().endswith(client_suffix.casefold()):
                candidate = candidate[: -len(client_suffix)].strip(" -")
            if candidate:
                return candidate
    return ""


def _project_from_filename(filename: str, proposal_number: str, client_name: str) -> str:
    if not filename:
        return ""
    candidate = Path(filename).stem.strip()
    if candidate.casefold().startswith(proposal_number.casefold()):
        candidate = candidate[len(proposal_number) :].strip(" -")
    client_suffix = f" - {client_name}" if client_name else ""
    if client_suffix and candidate.casefold().endswith(client_suffix.casefold()):
        candidate = candidate[: -len(client_suffix)].strip(" -")
    return candidate


def extract_uploaded_proposal(docx_bytes: bytes, filename: str = "") -> UploadedProposalDetails:
    try:
        document = Document(io.BytesIO(docx_bytes))
    except Exception as exc:
        raise ValueError("The uploaded file is not a readable Word proposal.") from exc

    paragraphs = document.paragraphs
    introduction_index = next(
        (
            index
            for index, paragraph in enumerate(paragraphs)
            if paragraph.text.strip() == "Introduction"
        ),
        None,
    )
    front_paragraphs = paragraphs[:introduction_index] if introduction_index is not None else paragraphs[:10]
    front_lines = [
        line
        for paragraph in front_paragraphs
        for line in _nonempty_lines(paragraph.text)
    ]
    proposal_line_index = next(
        (
            index
            for index, line in enumerate(front_lines)
            if "almor proposal no." in line.casefold()
        ),
        None,
    )
    if proposal_line_index is None:
        raise ValueError("Could not find 'Almor Proposal No.' in the uploaded Word proposal.")

    proposal_line = front_lines[proposal_line_index]
    number_match = PROPOSAL_NUMBER_PATTERN.search(proposal_line)
    proposal_number = number_match.group(0).upper() if number_match else ""
    date_text = re.split(
        r"almor proposal no\.?:?",
        proposal_line,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0]
    proposal_date = _iso_date(date_text.strip(" \t|-:"))

    attention_index = next(
        (
            index
            for index, line in enumerate(front_lines)
            if line.casefold().startswith("attention:")
        ),
        None,
    )
    re_index = next(
        (
            index
            for index, line in enumerate(front_lines)
            if re.match(r"^re\s*:", line, flags=re.IGNORECASE)
        ),
        len(front_lines),
    )

    client_end = min(attention_index if attention_index is not None else len(front_lines), re_index)
    client_lines = front_lines[proposal_line_index + 1 : client_end]
    client_name = client_lines[0] if client_lines else ""
    client_address = "\n".join(client_lines[1:])

    contact_name = ""
    contact_email = ""
    if attention_index is not None:
        attention = front_lines[attention_index]
        contact = attention.split(":", 1)[-1].strip()
        contact_parts = [part.strip() for part in contact.split("|", 1)]
        contact_name = contact_parts[0]
        if len(contact_parts) > 1:
            contact_email = contact_parts[1]

    project_name, project_location = _project_details(front_lines)
    if not project_name:
        project_name = _project_from_identifier(paragraphs, proposal_number, client_name)
    if not project_name:
        project_name = _project_from_filename(filename, proposal_number, client_name)

    budget = _proposal_budget(document)
    missing = [
        label
        for label, value in (
            ("proposal number", proposal_number),
            ("proposal date", proposal_date),
            ("client", client_name),
            ("estimated cost", budget),
        )
        if not value
    ]
    if missing:
        raise ValueError("Could not read these required proposal fields: " + ", ".join(missing) + ".")

    facts = ProposalFacts(
        proposal_number=proposal_number,
        proposal_date=proposal_date,
        client_name=client_name,
        client_address=client_address,
        contact_name=contact_name,
        contact_email=contact_email,
        project_name=project_name,
        project_location=project_location,
    )
    scope = f"{project_name} - geotechnical services described in {proposal_number}."
    return UploadedProposalDetails(
        facts=facts,
        work_authorization_scope=scope,
        budget=float(budget),
    )
