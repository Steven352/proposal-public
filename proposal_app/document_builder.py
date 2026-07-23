from __future__ import annotations

import io
import re
import zipfile
from copy import deepcopy
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table, _Row
from docx.text.paragraph import Paragraph
from lxml import etree

from .config import PREPARED_BY, PREPARED_BY_TITLE, REVIEWED_BY, REVIEWED_BY_TITLE, SECTION_ORDER
from .costs import format_money, normalize_cost_items
from .models import DraftContent, ParagraphBlock, ProposalFacts


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
XML_SPACE = "{http://www.w3.org/XML/1998/namespace}space"


def safe_filename(value: str) -> str:
    cleaned = re.sub(r'[<>:"/\\|?*]+', " ", value)
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" .")
    return cleaned or "Proposal"


def output_stem(facts: ProposalFacts) -> str:
    number = safe_filename(facts.proposal_number or "Proposal")
    project = safe_filename(facts.project_name or "Geotechnical Assessment")
    client = safe_filename(facts.client_name or "Client")
    return f"{number} {project} - {client}"


def parse_date(value: str) -> date:
    for parser in (
        lambda raw: date.fromisoformat(raw),
        lambda raw: datetime.strptime(raw, "%B %d, %Y").date(),
        lambda raw: datetime.strptime(raw, "%b %d, %Y").date(),
    ):
        try:
            return parser(value.strip())
        except (ValueError, TypeError):
            continue
    return date.today()


def format_date(value: str) -> str:
    parsed = parse_date(value)
    return f"{parsed.strftime('%B')} {parsed.day}, {parsed.year}"


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run._element.getparent().remove(run._element)
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        extra._element.getparent().remove(extra._element)


def paragraph_by_text(document: DocumentType, text: str) -> Paragraph:
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == text:
            return paragraph
    raise ValueError(f"Required heading not found: {text}")


def paragraphs_between(document: DocumentType, start: str, end: str | None) -> list[Paragraph]:
    paragraphs = document.paragraphs
    start_index = next(i for i, p in enumerate(paragraphs) if p.text.strip() == start)
    if end is None:
        end_index = len(paragraphs)
    else:
        end_index = next(
            i for i, p in enumerate(paragraphs[start_index + 1 :], start_index + 1)
            if p.text.strip() == end
        )
    return paragraphs[start_index + 1 : end_index]


def replace_xml_paragraph_text(paragraph_xml, text: str) -> None:
    text_nodes = paragraph_xml.xpath(".//w:t")
    if not text_nodes:
        return
    text_nodes[0].text = text
    text_nodes[0].set(XML_SPACE, "preserve")
    for node in text_nodes[1:]:
        node.text = ""


def clone_paragraph_before(anchor: Paragraph, prototype: Paragraph, text: str):
    clone = deepcopy(prototype._p)
    replace_xml_paragraph_text(clone, text)
    anchor._p.addprevious(clone)


def replace_section(
    document: DocumentType,
    start_heading: str,
    end_heading: str,
    blocks: list[ParagraphBlock],
) -> None:
    existing = paragraphs_between(document, start_heading, end_heading)
    if not existing:
        raise ValueError(f"No editable content found below {start_heading}")
    body_prototype = next((p for p in existing if p.style.name == "Body"), existing[0])
    list_prototype = next(
        (
            p
            for p in existing
            if p._p.pPr is not None and p._p.pPr.numPr is not None
        ),
        body_prototype,
    )
    for paragraph in existing:
        paragraph._element.getparent().remove(paragraph._element)
    anchor = paragraph_by_text(document, end_heading)
    for block in blocks:
        prototype = list_prototype if block.style == "list" else body_prototype
        clone_paragraph_before(anchor, prototype, block.text)


def update_front_matter(document: DocumentType, facts: ProposalFacts) -> dict[str, str]:
    introduction = paragraph_by_text(document, "Introduction")
    preceding = []
    for paragraph in document.paragraphs:
        if paragraph._p is introduction._p:
            break
        preceding.append(paragraph)
    if len(preceding) < 5:
        raise ValueError("Selected proposal does not have a compatible front matter layout.")

    old_number_match = re.search(r"P0\d{2}-[0-9A-Za-z]+", preceding[0].text)
    old_number = old_number_match.group(0) if old_number_match else ""
    old_client = preceding[1].text.splitlines()[0].strip() if len(preceding) > 1 else ""
    old_project = preceding[4].text.splitlines()[-1].strip() if len(preceding) > 4 else ""
    old_location = next(
        (
            line.strip().lstrip("\t")
            for paragraph in preceding[5:]
            for line in paragraph.text.splitlines()
            if line.strip()
        ),
        "",
    )

    set_paragraph_text(
        preceding[0],
        f"{format_date(facts.proposal_date)}\tAlmor Proposal No.: {facts.proposal_number}",
    )
    client_block = "\n".join(
        value for value in [facts.client_name, facts.client_address] if value.strip()
    )
    set_paragraph_text(preceding[1], client_block)
    if len(preceding) > 2:
        set_paragraph_text(preceding[2], "")
    attention_parts = [facts.contact_name]
    if facts.contact_title:
        attention_parts[0] = f"{facts.contact_name}, {facts.contact_title}" if facts.contact_name else facts.contact_title
    attention = f"Attention: {attention_parts[0]}" if attention_parts[0] else ""
    if facts.contact_email:
        attention = f"{attention} | {facts.contact_email}" if attention else facts.contact_email
    if len(preceding) > 3:
        set_paragraph_text(preceding[3], attention)
    set_paragraph_text(
        preceding[4],
        f"Re:\tRequest for Proposal for Geotechnical Services\n{facts.project_name}",
    )
    location_lines = [line.strip() for line in facts.project_location.splitlines() if line.strip()]
    for index, paragraph in enumerate(preceding[5:]):
        set_paragraph_text(paragraph, location_lines[index] if index < len(location_lines) else "")

    return {
        "old_number": old_number,
        "old_client": old_client,
        "old_project": old_project,
        "old_location": old_location,
    }


def update_cost_table(table: Table, facts: ProposalFacts) -> float:
    summary = normalize_cost_items(facts.cost_items)
    if summary.warnings:
        raise ValueError(" ".join(summary.warnings))
    if not summary.items:
        raise ValueError("A non-empty cost table is required.")

    rows = table.rows
    if len(rows) < 4 or len(table.columns) != 6:
        raise ValueError("Selected proposal has an incompatible cost table.")
    section_template = deepcopy(rows[1]._tr)
    item_template = deepcopy(rows[2]._tr)
    subtotal_template = deepcopy(rows[3]._tr)
    final_template = deepcopy(rows[-1]._tr)
    for row in list(table.rows[1:]):
        table._tbl.remove(row._tr)

    by_section = {section: [] for section in SECTION_ORDER}
    for item in summary.items:
        by_section[item.section].append(item)

    for section in SECTION_ORDER:
        items = by_section[section]
        if not items:
            continue
        table._tbl.append(deepcopy(section_template))
        set_row_values(table.rows[-1], [section, "", "", "", "", ""])
        for item in items:
            table._tbl.append(deepcopy(item_template))
            estimate = "" if item.estimate is None else f"{item.estimate:g}"
            set_row_values(
                table.rows[-1],
                [
                    item.item,
                    item.description,
                    item.unit,
                    estimate,
                    format_money(item.rate),
                    format_money(item.total),
                ],
            )
        table._tbl.append(deepcopy(subtotal_template))
        set_row_values(
            table.rows[-1],
            ["", "", "", "", "Subtotal", format_money(summary.section_totals[section])],
        )

    table._tbl.append(deepcopy(final_template))
    set_row_values(
        table.rows[-1],
        ["", "", "", "", "Estimated Cost", format_money(summary.final_total)],
    )
    return summary.final_total


def set_row_values(row: _Row, values: list[str]) -> None:
    for cell, value in zip(row.cells, values, strict=True):
        set_cell_text(cell, value)


def update_cost_section(
    document: DocumentType,
    draft: DraftContent,
    facts: ProposalFacts,
    final_total: float,
) -> None:
    blocks = list(document.element.body.iterchildren())
    cost_heading = paragraph_by_text(document, "Cost of Geotechnical Services")._p
    terms_heading = paragraph_by_text(document, "Terms and Conditions")._p
    start = blocks.index(cost_heading)
    end = blocks.index(terms_heading)
    between = blocks[start + 1 : end]
    paragraphs = [Paragraph(node, document) for node in between if node.tag.endswith("}p")]
    if len(paragraphs) < 2:
        raise ValueError("Selected proposal cost section is incompatible.")
    set_paragraph_text(paragraphs[0], draft.cost_intro)
    total_text = (
        "The total estimated geotechnical engineering project services cost for this work is "
        f"{format_money(final_total)}. The estimated cost includes all professional service fees, "
        "surcharges, associated fees and is valid for a period of 60 days from the date on the "
        "proposal. The estimated cost is based on the scope of work and assumptions described in "
        "this document and is exclusive of GST and any other applicable taxes. Invoices would be "
        "forwarded to the client monthly."
    )
    set_paragraph_text(paragraphs[-1], total_text)
    for paragraph in paragraphs[1:-1]:
        paragraph._element.getparent().remove(paragraph._element)


def update_closure(document: DocumentType, draft: DraftContent, facts: ProposalFacts) -> None:
    paragraphs = paragraphs_between(document, "Closure", None)
    if len(paragraphs) < 8:
        raise ValueError("Selected proposal closure/signature block is incompatible.")
    values = [
        draft.closure_paragraph_1,
        draft.closure_paragraph_2,
        "Respectfully Submitted,",
        "ALMOR TESTING SERVICES LTD.",
        "Prepared by\t\t\t\t\t\t\tReviewed By",
        f"{PREPARED_BY}\t\t\t\t \t           {REVIEWED_BY}",
        f"{PREPARED_BY_TITLE}\t\t\t\t                       {REVIEWED_BY_TITLE}",
        output_stem(facts),
    ]
    for paragraph, value in zip(paragraphs[:8], values, strict=True):
        set_paragraph_text(paragraph, value)
    for paragraph in paragraphs[8:]:
        paragraph._element.getparent().remove(paragraph._element)


def patch_package_text(docx_bytes: bytes, replacements: dict[str, str]) -> bytes:
    source = io.BytesIO(docx_bytes)
    target = io.BytesIO()
    with zipfile.ZipFile(source) as zin, zipfile.ZipFile(target, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                try:
                    root = etree.fromstring(data)
                    changed = False
                    for paragraph in root.xpath(".//w:p", namespaces={"w": W_NS}):
                        text_nodes = paragraph.xpath(".//w:t", namespaces={"w": W_NS})
                        if not text_nodes:
                            continue
                        for old, new in replacements.items():
                            if old and replace_text_across_nodes(text_nodes, old, new):
                                changed = True
                    if changed:
                        data = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone="yes")
                except etree.XMLSyntaxError:
                    pass
            zout.writestr(item, data)
    return target.getvalue()


def replace_text_across_nodes(text_nodes: list, old: str, new: str) -> bool:
    """Replace text without collapsing the template's separately formatted runs."""
    changed = False
    while True:
        values = [node.text or "" for node in text_nodes]
        combined = "".join(values)
        start = combined.rfind(old)
        if start < 0:
            return changed
        end = start + len(old)

        offsets = []
        cursor = 0
        for value in values:
            offsets.append((cursor, cursor + len(value)))
            cursor += len(value)

        start_index = next(
            index for index, (_, node_end) in enumerate(offsets) if start < node_end
        )
        end_index = next(
            index for index, (_, node_end) in enumerate(offsets) if end <= node_end
        )
        start_offset = start - offsets[start_index][0]
        end_offset = end - offsets[end_index][0]

        prefix = values[start_index][:start_offset]
        suffix = values[end_index][end_offset:]
        text_nodes[start_index].text = prefix + new + (suffix if start_index == end_index else "")
        text_nodes[start_index].set(XML_SPACE, "preserve")
        if start_index != end_index:
            for index in range(start_index + 1, end_index):
                text_nodes[index].text = ""
            text_nodes[end_index].text = suffix
            if suffix:
                text_nodes[end_index].set(XML_SPACE, "preserve")
        changed = True


def find_header_location_replacements(
    docx_bytes: bytes,
    old_location: str,
    old_client: str,
    new_location: str,
) -> dict[str, str]:
    """Locate the template's visible header location slot without rebuilding its text box."""
    if not old_location or not new_location:
        return {}
    replacements: dict[str, str] = {}
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
        for name in archive.namelist():
            if not name.startswith("word/header") or not name.endswith(".xml"):
                continue
            root = etree.fromstring(archive.read(name))
            for paragraph in root.xpath(
                ".//w:p[not(.//w:p)]",
                namespaces={"w": W_NS},
            ):
                text = "".join(
                    paragraph.xpath(".//w:t/text()", namespaces={"w": W_NS})
                ).strip()
                if (
                    old_location.casefold() in text.casefold()
                    and not (old_client and old_client.casefold() in text.casefold())
                    and "Almor Proposal No." not in text
                ):
                    replacements[text] = new_location
    return replacements


def build_docx(
    template_path: Path,
    facts: ProposalFacts,
    draft: DraftContent,
) -> tuple[bytes, str]:
    document = Document(template_path)
    old_values = update_front_matter(document, facts)

    replace_section(
        document,
        "Introduction",
        "Proposed Project Personnel" if any(
            p.text.strip() == "Proposed Project Personnel" for p in document.paragraphs
        ) else "Scope of Work and Cost Summary",
        [ParagraphBlock(text=draft.introduction, style="body")],
    )

    field_blocks = [ParagraphBlock(text=draft.field_program_intro, style="body")]
    field_blocks.extend(draft.field_program_paragraphs)
    replace_section(
        document,
        "Geotechnical Field Program",
        "Cost of Geotechnical Services",
        field_blocks,
    )

    cost_table = next((table for table in document.tables if len(table.columns) == 6), None)
    if cost_table is None:
        raise ValueError("Selected proposal does not contain a compatible six-column cost table.")
    final_total = update_cost_table(cost_table, facts)
    update_cost_section(document, draft, facts, final_total)

    terms_paragraphs = paragraphs_between(document, "Terms and Conditions", "Closure")
    if not terms_paragraphs:
        raise ValueError("Selected proposal terms section is incompatible.")
    set_paragraph_text(terms_paragraphs[0], draft.terms_and_conditions)
    for paragraph in terms_paragraphs[1:]:
        paragraph._element.getparent().remove(paragraph._element)
    update_closure(document, draft, facts)

    buffer = io.BytesIO()
    document.save(buffer)
    saved_bytes = buffer.getvalue()
    header_location = ", ".join(
        line.strip() for line in facts.project_location.splitlines() if line.strip()
    )
    replacements = {
        old_values["old_number"]: facts.proposal_number,
        old_values["old_client"]: facts.client_name,
        old_values["old_project"]: facts.project_name,
        "Steven Lai, E.I.T.": PREPARED_BY,
    }
    replacements.update(
        find_header_location_replacements(
            saved_bytes,
            old_values["old_location"],
            old_values["old_client"],
            header_location,
        )
    )
    result = patch_package_text(saved_bytes, replacements)
    return result, output_stem(facts) + ".docx"
